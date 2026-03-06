import os
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any
import docx
import fitz  # PyMuPDF
from pathlib import Path

DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _first_non_empty_paragraph(doc: Any) -> str:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            return text
    return ""


def _extract_abstract_and_keywords(full_text: str) -> tuple[str, str]:
    lines = [line.strip() for line in full_text.splitlines()]
    abstract = ""
    keywords = ""

    abstract_re = re.compile(r"^[【\[]?\s*摘要\s*[】\]]?\s*[:：]?\s*(.*)$")
    keywords_re = re.compile(r"^[【\[]?\s*(关键词|关键字)\s*[】\]]?\s*[:：]?\s*(.*)$")

    for i, line in enumerate(lines):
        if not line:
            continue
        if not abstract:
            m = abstract_re.match(line)
            if m:
                value = (m.group(1) or "").strip()
                if value:
                    abstract = value
                else:
                    for j in range(i + 1, len(lines)):
                        nxt = lines[j]
                        if not nxt:
                            continue
                        if keywords_re.match(nxt):
                            break
                        abstract = nxt
                        break
        if not keywords:
            m = keywords_re.match(line)
            if m:
                keywords = (m.group(2) or "").strip()

    if not abstract:
        m = re.search(
            r"[【\[]?\s*摘要\s*[】\]]?\s*[:：]?\s*(.+?)(?:\n\s*[【\[]?\s*(?:关键词|关键字)\s*[】\]]?\s*[:：]|\Z)",
            full_text,
            flags=re.S,
        )
        if m:
            abstract = re.sub(r"\s+", " ", m.group(1)).strip()

    if not keywords:
        m = re.search(r"[【\[]?\s*(?:关键词|关键字)\s*[】\]]?\s*[:：]?\s*(.+)", full_text)
        if m:
            keywords = m.group(1).splitlines()[0].strip()

    return abstract, keywords


def _extract_notes_from_docx_xml(file_path: Path, tag_name: str) -> list[str]:
    xml_path = f"word/{tag_name}.xml"
    notes: list[str] = []
    if not file_path.exists():
        return notes

    try:
        with zipfile.ZipFile(file_path) as zf:
            if xml_path not in zf.namelist():
                return notes
            xml_data = zf.read(xml_path)
    except Exception:
        return notes

    try:
        root = ET.fromstring(xml_data)
    except ET.ParseError:
        return notes

    node_name = "footnote" if tag_name == "footnotes" else "endnote"
    for note_node in root.findall(f"w:{node_name}", DOCX_NS):
        note_type = note_node.attrib.get(f"{{{DOCX_NS['w']}}}type")
        if note_type in {"separator", "continuationSeparator", "continuationNotice"}:
            continue
        note_id = note_node.attrib.get(f"{{{DOCX_NS['w']}}}id")
        text_chunks = [
            (t.text or "").strip()
            for t in note_node.findall(".//w:t", DOCX_NS)
            if (t.text or "").strip()
        ]
        if not text_chunks:
            continue
        note_text = " ".join(text_chunks)
        notes.append(f"{note_id}: {note_text}" if note_id else note_text)
    return notes


def parse_docx(file_path: Path) -> Dict[str, Any]:
    """解析 Word 文档。"""
    doc = docx.Document(file_path)
    full_text = []
    paragraphs = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
        # 简单识别标题（基于样式或加粗，初版简单处理）
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        if style_name.startswith('Heading') or any(run.bold for run in para.runs if run.text.strip()):
            paragraphs.append({"text": para.text, "is_header": True})
        else:
            paragraphs.append({"text": para.text, "is_header": False})

    text_str = "\n".join(full_text)
    abstract, keywords = _extract_abstract_and_keywords(text_str)
    footnotes = _extract_notes_from_docx_xml(file_path, "footnotes")
    endnotes = _extract_notes_from_docx_xml(file_path, "endnotes")
    notes = footnotes + endnotes

    return {
        "title": _first_non_empty_paragraph(doc),
        "abstract": abstract,
        "keywords": keywords,
        "body_text": text_str,
        "body_structure": paragraphs,
        "footnotes_raw": notes,
        "references_raw": [],
        "author_info": {},
        "word_count": len(re.sub(r"\s+", "", text_str))
    }

def parse_pdf(file_path: Path) -> Dict[str, Any]:
    """解析 PDF 文档。"""
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    
    # PDF 解析相对复杂，初版仅提取全文
    return {
        "title": os.path.basename(file_path),
        "abstract": "",
        "keywords": "",
        "body_text": text,
        "body_structure": [],
        "footnotes_raw": [],
        "references_raw": [],
        "author_info": {},
        "word_count": len(text)
    }

def parse_manuscript(file_path: str) -> Dict[str, Any]:
    """根据文件扩展名选择解析器。"""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == ".docx":
        return parse_docx(path)
    elif ext == ".pdf":
        return parse_pdf(path)
    else:
        # 不支持的格式，返回最小信息
        return {
            "title": path.name,
            "abstract": "",
            "keywords": "",
            "body_text": "",
            "body_structure": [],
            "footnotes_raw": [],
            "references_raw": [],
            "author_info": {},
            "word_count": 0
        }
